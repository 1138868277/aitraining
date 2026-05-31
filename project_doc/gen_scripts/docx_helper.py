"""
python-docx 共享工具模块
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def create_document(title: str):
    """创建文档并设置基本样式"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc

def add_cover(doc, title, subtitle, version="V1.0", date="2026-05-30"):
    """添加封面"""
    # 空行占位
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    # 副标题
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    for _ in range(6):
        doc.add_paragraph()

    # 版本信息
    info_items = [("版  本", version), ("日  期", date)]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(14)
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    doc.add_page_break()

def add_heading1(doc, text):
    """添加一级标题"""
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading2(doc, text):
    """添加二级标题"""
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading3(doc, text):
    """添加三级标题"""
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(doc, text, bold=False, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)
    run.font.bold = bold
    return p

def add_bullet(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    for run in p.runs:
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table

def add_toc(doc):
    """添加目录页"""
    add_heading1(doc, "目录")
    p = doc.add_paragraph("（生成后请在 Word 中右键点击此处 → 更新域 → 更新整个目录）")
    p.paragraph_format.first_line_indent = Pt(22)
    doc.add_page_break()

def save_doc(doc, filename):
    """保存文档"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"文档已生成：{filepath}")
    return filepath
